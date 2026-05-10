import base64
import logging
from fastapi import APIRouter, UploadFile, File
from fastapi.responses import JSONResponse

from app.core.config import settings

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/upload", tags=["upload"])


def _extract_pdf_text(file_bytes: bytes) -> tuple[str | None, str]:
    """尝试提取 PDF 文本内容

    Returns:
        tuple: (提取的文本或None, 错误信息)
    """
    try:
        import fitz
        import io
    except ImportError:
        return None, "PyMuPDF库未安装，无法解析PDF文件"

    try:
        doc = fitz.open(stream=io.BytesIO(file_bytes), filetype="pdf")
        if doc.is_closed:
            return None, "PDF文件格式错误，无法打开"
        text = ""
        for page_num, page in enumerate(doc):
            try:
                page_text = page.get_text()
                text += page_text
            except Exception as e:
                logger.warning(f"[PDF] 第{page_num + 1}页提取失败: {e}")
                continue
        doc.close()

        if not text.strip():
            return None, "PDF文件中没有可提取的文本内容（可能是扫描件或图片型PDF）"
        return text.strip(), ""
    except Exception as e:
        error_msg = f"PDF文件解析失败: {str(e)}"
        logger.warning(f"[UploadForward] {error_msg}")
        return None, error_msg


def _extract_docx_text(file_bytes: bytes) -> tuple[str | None, str]:
    """尝试提取 Word 文档文本内容

    Returns:
        tuple: (提取的文本或None, 错误信息)
    """
    try:
        import docx
        import io
    except ImportError:
        return None, "python-docx库未安装，无法解析Word文档"

    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        text = "\n".join(paragraphs)
        if not text:
            return None, "Word文档中没有可提取的文本内容"
        return text, ""
    except Exception as e:
        error_msg = f"Word文档解析失败: {str(e)}"
        logger.warning(f"[UploadForward] {error_msg}")
        return None, error_msg


@router.post("/forward")
async def forward_file(file: UploadFile = File(...)):
    try:
        file_bytes = await file.read()
    except Exception as e:
        logger.warning(f"[UploadForward] 文件读取失败: {e}")
        return JSONResponse(
            status_code=400,
            content={"status": "error", "message": "文件读取失败，请重新上传"},
        )

    if len(file_bytes) > settings.FILE_MAX_SIZE:
        return JSONResponse(
            status_code=400,
            content={"status": "error", "message": f"文件大小超过限制 ({settings.FILE_MAX_SIZE // 1024 // 1024}MB)"},
        )

    try:
        filename = file.filename or "unknown"
        content_type = file.content_type or ""
        ext = filename.split('.')[-1].lower() if '.' in filename else ''

        if not ext and content_type:
            ext_map = {
                "image/jpeg": "jpg",
                "image/png": "png",
                "image/gif": "gif",
                "image/bmp": "bmp",
                "image/webp": "webp",
                "application/pdf": "pdf",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
            }
            ext = ext_map.get(content_type, "")

        image_extensions = {'jpg', 'jpeg', 'png', 'gif', 'bmp', 'webp'}
        if ext in image_extensions:
            b64 = base64.b64encode(file_bytes).decode('utf-8')
            mime = f"image/{ext if ext != 'jpg' else 'jpeg'}"
            content = f"data:{mime};base64,{b64}"
            return {
                "status": "success",
                "content": content,
                "type": "image",
                "filename": filename,
            }

        if ext == 'pdf':
            text, error = _extract_pdf_text(file_bytes)
            if text:
                return {
                    "status": "success",
                    "content": text,
                    "type": "text",
                    "filename": filename,
                }
            return JSONResponse(
                status_code=400,
                content={"status": "error", "message": error or "无法提取PDF文本内容"},
            )

        if ext in ('docx', 'doc'):
            text, error = _extract_docx_text(file_bytes)
            if text:
                return {
                    "status": "success",
                    "content": text,
                    "type": "text",
                    "filename": filename,
                }
            return JSONResponse(
                status_code=400,
                content={"status": "error", "message": error or "无法提取Word文档文本内容"},
            )

        text_extensions = {'txt', 'md', 'csv', 'json', 'xml', 'html', 'css', 'js', 'py', 'java', 'cpp', 'c', 'h', 'go', 'rs', 'ts', 'sql', 'yaml', 'yml'}
        if ext in text_extensions:
            try:
                text = file_bytes.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    text = file_bytes.decode('gbk')
                except UnicodeDecodeError:
                    text = file_bytes.decode('utf-8', errors='ignore')
            return {
                "status": "success",
                "content": text,
                "type": "text",
                "filename": filename,
            }

        return JSONResponse(
            status_code=400,
            content={"status": "error", "message": f"不支持的文件类型: .{ext}"},
        )

    except Exception as e:
        logger.warning(f"[UploadForward] 文件处理失败: {e}")
        return JSONResponse(
            status_code=500,
            content={"status": "error", "message": "文件上传处理失败，请稍后重试"},
        )
